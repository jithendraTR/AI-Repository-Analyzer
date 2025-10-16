"""
Base Analyzer Class
Common functionality for all repository analyzers
"""

import os
import git
import streamlit as st
from abc import ABC, abstractmethod
from typing import Dict, List, Any, Optional
from pathlib import Path
import subprocess
import json
import pandas as pd
from datetime import datetime, timedelta
import threading
import time
from utils.ai_client import OpenArenaClient

class CancellationToken:
    """Thread-safe cancellation token for stopping long-running operations"""
    
    def __init__(self, token_id: str):
        self.token_id = token_id
        self._cancelled = threading.Event()
        self._start_time = time.time()
    
    def cancel(self):
        """Cancel the operation"""
        self._cancelled.set()
        # Update session state
        if f"cancelled_{self.token_id}" not in st.session_state:
            st.session_state[f"cancelled_{self.token_id}"] = True
    
    def is_cancelled(self) -> bool:
        """Check if the operation has been cancelled"""
        # Check both the event and session state
        session_cancelled = st.session_state.get(f"cancelled_{self.token_id}", False)
        return self._cancelled.is_set() or session_cancelled
    
    def check_cancellation(self):
        """Raise an exception if cancelled"""
        if self.is_cancelled():
            raise OperationCancelledException(f"Operation {self.token_id} was cancelled")
    
    def get_elapsed_time(self) -> float:
        """Get elapsed time since token creation"""
        return time.time() - self._start_time
    
    def cleanup(self):
        """Clean up session state"""
        keys_to_remove = [
            f"cancelled_{self.token_id}",
            f"running_{self.token_id}",
            f"progress_{self.token_id}"
        ]
        for key in keys_to_remove:
            if key in st.session_state:
                del st.session_state[key]

class OperationCancelledException(Exception):
    """Exception raised when an operation is cancelled"""
    pass

class BaseAnalyzer(ABC):
    """Base class for all repository analyzers"""
    
    def __init__(self, repo_path: str):
        # Validate repository path
        if not repo_path or not repo_path.strip():
            raise ValueError("Repository path cannot be empty")
        
        if repo_path in ["/path/to/your/repo", "C:\\path\\to\\your\\repo"]:
            raise ValueError("Please provide a valid repository path, not the placeholder text")
        
        if not os.path.exists(repo_path):
            raise ValueError(f"Repository path does not exist: {repo_path}")
        
        self.repo_path = Path(repo_path)
        self.ai_client = OpenArenaClient()
        
        # Initialize git repository if it exists
        try:
            self.repo = git.Repo(repo_path)
        except git.exc.InvalidGitRepositoryError:
            self.repo = None
            st.warning("Not a git repository - some features may be limited")
    
    @abstractmethod
    def analyze(self, token=None, progress_callback=None) -> Dict[str, Any]:
        """
        Perform the specific analysis - must be implemented by subclasses
        
        Args:
            token: Optional cancellation token for stopping long-running operations
            progress_callback: Optional callback function for progress updates
        """
        pass
    
    @abstractmethod
    def render(self):
        """Render the analysis results in Streamlit - must be implemented by subclasses"""
        pass
    
    def render_with_header(self, analysis_type: str, title: str, description: str = None):
        """
        Render analysis with a single header - use this in subclass render methods
        
        Args:
            analysis_type: Type of analysis for caching/controls
            title: Display title for the analysis
            description: Optional description text
        """
        # Display header only once
        st.header(f"🔍 {title}")
        if description:
            st.markdown(description)
        st.markdown("---")
        
        # Add rerun button
        self.add_rerun_button(analysis_type)
        
        # Display any parallel AI insights first
        insights_displayed = self.display_parallel_ai_insights(analysis_type)
        if insights_displayed:
            st.markdown("---")
    
    def render_analysis_section(self, analysis_type: str, loading_message: str):
        """
        Render analysis section without header - use this for the main analysis content
        
        Args:
            analysis_type: Type of analysis
            loading_message: Message to show while loading
            
        Returns:
            Analysis results or None if analysis should not run
        """
        return self.get_analysis_with_control(analysis_type, loading_message)

    def get_file_list(self, extensions: List[str] = None) -> List[Path]:
        """
        Get list of files in the repository
        
        Args:
            extensions: List of file extensions to filter by (e.g., ['.py', '.js'])
            
        Returns:
            List of file paths
        """
        files = []
        for root, dirs, filenames in os.walk(self.repo_path):
            # Skip common directories that shouldn't be analyzed
            dirs[:] = [d for d in dirs if not d.startswith('.') and d not in 
                      ['node_modules', '__pycache__', 'venv', 'env', 'build', 'dist']]
            
            for filename in filenames:
                file_path = Path(root) / filename
                if extensions:
                    if file_path.suffix.lower() in extensions:
                        files.append(file_path)
                else:
                    files.append(file_path)
        
        return files
    
    def read_file_content(self, file_path: Path) -> Optional[str]:
        """
        Read content of a file safely
        
        Args:
            file_path: Path to the file
            
        Returns:
            File content or None if error
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except PermissionError:
            # Silently skip files with permission issues (common in temp directories)
            return None
        except (UnicodeDecodeError, IsADirectoryError):
            # Silently skip binary files or directories
            return None
        except Exception as e:
            # Only show warning for unexpected errors, not for temp/clone directories
            if "temp" not in str(file_path).lower() and "clone" not in str(file_path).lower():
                st.warning(f"Could not read file {file_path}: {str(e)}")
            return None
    
    def get_git_history(self, file_path: str = None, max_commits: int = 100) -> List[Dict]:
        """
        Get git commit history with time frame filtering applied
        
        Args:
            file_path: Specific file to get history for (optional)
            max_commits: Maximum number of commits to retrieve
            
        Returns:
            List of commit information dictionaries filtered by time frame
        """
        if not self.repo:
            # Check if we should show an error for no commits
            error_info = self.check_commit_filter_error()
            if not error_info:
                # Set error for no git repository
                selected_time_frame = st.session_state.get('selected_time_frame', 'all')
                if selected_time_frame != 'all':
                    time_frame_display = {
                        '1_year': 'last 1 year',
                        '2_years': 'last 2 years',
                        '3_years': 'last 3 years', 
                        '5_years': 'last 5 years'
                    }.get(selected_time_frame, selected_time_frame)
                    
                    st.session_state['commit_filter_error'] = {
                        'message': f"No git repository found. Unable to analyze commits for the {time_frame_display}.",
                        'total_commits': 0,
                        'selected_period': time_frame_display,
                        'has_commits': False
                    }
            return []
        
        try:
            commits = []
            commit_iter = self.repo.iter_commits(paths=file_path, max_count=max_commits)
            
            for commit in commit_iter:
                commits.append({
                    'hash': commit.hexsha,
                    'author': commit.author.name,
                    'email': commit.author.email,
                    'date': commit.committed_datetime,
                    'message': commit.message.strip(),
                    'files_changed': len(commit.stats.files)
                })
            
            # Apply time frame filtering
            return self.filter_commits_by_time_frame(commits)
        except Exception as e:
            # Completely suppress git history warnings for cloned repositories
            # Only log if it's a local repository that should have working git
            return []
    
    def filter_commits_by_time_frame(self, commits: List[Dict]) -> List[Dict]:
        """
        Filter commits based on selected time frame from session state with timezone-aware handling
        
        Args:
            commits: List of commit dictionaries with 'date' field
            
        Returns:
            Filtered list of commits
        """
        # Get selected time frame from session state
        selected_time_frame = st.session_state.get('selected_time_frame', 'all')
        
        # If 'all' is selected, return all commits
        if selected_time_frame == 'all':
            # Clear any previous error when showing all commits
            if 'commit_filter_error' in st.session_state:
                del st.session_state['commit_filter_error']
            return commits
        
        # Calculate cutoff date based on selection using UTC for consistency
        from datetime import timezone
        now_utc = datetime.now(timezone.utc)
        cutoff_date_utc = None
        
        if selected_time_frame == '1_year':
            cutoff_date_utc = now_utc - timedelta(days=365)
        elif selected_time_frame == '2_years':
            cutoff_date_utc = now_utc - timedelta(days=2*365)
        elif selected_time_frame == '3_years':
            cutoff_date_utc = now_utc - timedelta(days=3*365)
        elif selected_time_frame == '5_years':
            cutoff_date_utc = now_utc - timedelta(days=5*365)
        else:
            # Default to all commits if unknown selection
            if 'commit_filter_error' in st.session_state:
                del st.session_state['commit_filter_error']
            return commits
        
        # Filter commits based on cutoff date with proper timezone handling
        filtered_commits = []
        for commit in commits:
            commit_date = commit.get('date')
            if commit_date:
                try:
                    # Normalize commit date to UTC for consistent comparison
                    if hasattr(commit_date, 'tzinfo'):
                        if commit_date.tzinfo is not None:
                            # Convert timezone-aware datetime to UTC
                            commit_date_utc = commit_date.astimezone(timezone.utc)
                        else:
                            # Treat naive datetime as UTC (common for Git commits)
                            commit_date_utc = commit_date.replace(tzinfo=timezone.utc)
                    else:
                        # Handle other datetime formats by converting to string then parsing
                        if isinstance(commit_date, str):
                            try:
                                parsed_date = pd.to_datetime(commit_date)
                                if parsed_date.tzinfo is not None:
                                    commit_date_utc = parsed_date.astimezone(timezone.utc)
                                else:
                                    commit_date_utc = parsed_date.replace(tzinfo=timezone.utc)
                            except:
                                # If parsing fails, skip this commit from filtering
                                filtered_commits.append(commit)
                                continue
                        else:
                            # If not a recognizable date format, include the commit
                            filtered_commits.append(commit)
                            continue
                    
                    # Compare UTC-normalized dates
                    if commit_date_utc >= cutoff_date_utc:
                        filtered_commits.append(commit)
                        
                except Exception as e:
                    # On any timezone conversion error, include the commit to avoid breaking
                    # This ensures backwards compatibility
                    filtered_commits.append(commit)
                    continue
            else:
                # If no date field, include the commit to maintain existing behavior
                filtered_commits.append(commit)
        
        # Check if filtering resulted in empty commits and provide appropriate error handling
        if not filtered_commits and commits and selected_time_frame != 'all':
            # Store error information in session state for analyzers to check
            time_frame_display = {
                '1_year': 'last 1 year',
                '2_years': 'last 2 years', 
                '3_years': 'last 3 years',
                '5_years': 'last 5 years'
            }.get(selected_time_frame, selected_time_frame)
            
            st.session_state['commit_filter_error'] = {
                'message': f"No commits found for the {time_frame_display}. The repository has {len(commits)} total commits, but none fall within the selected time period.",
                'total_commits': len(commits),
                'selected_period': time_frame_display,
                'has_commits': len(commits) > 0
            }
        else:
            # Clear any previous error when commits are found
            if 'commit_filter_error' in st.session_state:
                del st.session_state['commit_filter_error']
        
        return filtered_commits
    
    def check_commit_filter_error(self) -> Dict[str, Any]:
        """
        Check if there's a commit filtering error and return error details
        
        Returns:
            Dictionary with error details or empty dict if no error
        """
        return st.session_state.get('commit_filter_error', {})
    
    def display_commit_filter_error(self) -> bool:
        """
        Display commit filter error if present
        
        Returns:
            True if error was displayed, False otherwise
        """
        error_info = self.check_commit_filter_error()
        if error_info:
            st.error(f"⚠️ {error_info['message']}")
            
            # Show helpful suggestions
            if error_info['has_commits']:
                st.info("💡 **Suggestions:**\n"
                       "- Try selecting 'All commits' to see the full repository history\n"
                       "- Choose a longer time period if the repository is older\n"
                       "- The repository may have been inactive during the selected period")
            else:
                st.warning("This repository has no commit history available.")
            
            return True
        return False
    
    def get_file_contributors(self, file_path: str) -> Dict[str, int]:
        """
        Get contributors to a specific file with their commit counts
        
        Args:
            file_path: Path to the file relative to repo root
            
        Returns:
            Dictionary of author -> commit count
        """
        if not self.repo:
            return {}
        
        try:
            contributors = {}
            commits = list(self.repo.iter_commits(paths=file_path))
            
            for commit in commits:
                author = commit.author.name
                contributors[author] = contributors.get(author, 0) + 1
            
            return contributors
        except Exception as e:
            st.warning(f"Could not get contributors for {file_path}: {str(e)}")
            return {}
    
    def find_files_by_pattern(self, pattern: str) -> List[Path]:
        """
        Find files matching a specific pattern
        
        Args:
            pattern: Glob pattern to match
            
        Returns:
            List of matching file paths
        """
        return list(self.repo_path.glob(pattern))
    
    def get_project_structure(self, max_depth: int = 3) -> str:
        """
        Get a string representation of the project structure
        
        Args:
            max_depth: Maximum directory depth to traverse
            
        Returns:
            String representation of project structure
        """
        structure = []
        
        def add_to_structure(path: Path, depth: int, prefix: str = ""):
            if depth > max_depth:
                return
            
            items = sorted([p for p in path.iterdir() if not p.name.startswith('.')])
            
            for i, item in enumerate(items):
                is_last = i == len(items) - 1
                current_prefix = "└── " if is_last else "├── "
                structure.append(f"{prefix}{current_prefix}{item.name}")
                
                if item.is_dir() and item.name not in ['node_modules', '__pycache__', 'venv', 'env']:
                    next_prefix = prefix + ("    " if is_last else "│   ")
                    add_to_structure(item, depth + 1, next_prefix)
        
        structure.append(self.repo_path.name)
        add_to_structure(self.repo_path, 0)
        
        return "\n".join(structure)
    
    def run_command(self, command: List[str]) -> Optional[str]:
        """
        Run a command in the repository directory
        
        Args:
            command: Command to run as list of strings
            
        Returns:
            Command output or None if error
        """
        try:
            result = subprocess.run(
                command,
                cwd=self.repo_path,
                capture_output=True,
                text=True,
                timeout=30
            )
            return result.stdout if result.returncode == 0 else None
        except Exception as e:
            st.warning(f"Command failed: {' '.join(command)} - {str(e)}")
            return None
    
    def cache_analysis(self, analysis_type: str, data: Any):
        """
        Cache analysis results
        
        Args:
            analysis_type: Type of analysis
            data: Data to cache
        """
        cache_key = f"{analysis_type}_{hash(str(self.repo_path))}"
        st.session_state[cache_key] = data
    
    def get_cached_analysis(self, analysis_type: str) -> Optional[Any]:
        """
        Get cached analysis results
        
        Args:
            analysis_type: Type of analysis
            
        Returns:
            Cached data or None if not found
        """
        cache_key = f"{analysis_type}_{hash(str(self.repo_path))}"
        return st.session_state.get(cache_key)
    
    def clear_cache(self, analysis_type: str = None):
        """
        Clear cached analysis results
        
        Args:
            analysis_type: Specific analysis type to clear, or None to clear all
        """
        if analysis_type:
            cache_key = f"{analysis_type}_{hash(str(self.repo_path))}"
            if cache_key in st.session_state:
                del st.session_state[cache_key]
        else:
            # Clear all cache for this repository
            repo_hash = hash(str(self.repo_path))
            keys_to_delete = [key for key in st.session_state.keys() 
                            if key.endswith(f"_{repo_hash}")]
            for key in keys_to_delete:
                del st.session_state[key]
    
    def create_cancellation_token(self, operation_name: str) -> CancellationToken:
        """
        Create a new cancellation token for an operation
        
        Args:
            operation_name: Name of the operation
            
        Returns:
            CancellationToken instance
        """
        token_id = f"{operation_name}_{hash(str(self.repo_path))}_{int(time.time())}"
        token = CancellationToken(token_id)
        
        # Mark operation as running
        st.session_state[f"running_{token.token_id}"] = True
        
        return token
    
    def display_cancellable_operation(self, token: CancellationToken, message: str, progress: float = None):
        """
        Display a cancellable operation with stop button
        
        Args:
            token: Cancellation token
            message: Status message
            progress: Progress value (0-100) if available
        """
        col1, col2 = st.columns([4, 1])
        
        with col1:
            if progress is not None:
                st.progress(progress / 100.0)
            st.text(f"{message} (Elapsed: {token.get_elapsed_time():.1f}s)")
        
        with col2:
            if st.button("🛑 Stop", key=f"stop_{token.token_id}"):
                token.cancel()
                st.warning("Stopping operation...")
                st.rerun()
    
    def get_file_list_cancellable(self, extensions: List[str] = None, token: CancellationToken = None) -> List[Path]:
        """
        Get list of files in the repository with cancellation support
        
        Args:
            extensions: List of file extensions to filter by
            token: Cancellation token
            
        Returns:
            List of file paths
        """
        files = []
        processed_dirs = 0
        
        for root, dirs, filenames in os.walk(self.repo_path):
            if token:
                token.check_cancellation()
            
            # Skip common directories that shouldn't be analyzed
            dirs[:] = [d for d in dirs if not d.startswith('.') and d not in 
                      ['node_modules', '__pycache__', 'venv', 'env', 'build', 'dist']]
            
            for filename in filenames:
                if token and processed_dirs % 10 == 0:  # Check every 10 files
                    token.check_cancellation()
                
                file_path = Path(root) / filename
                if extensions:
                    if file_path.suffix.lower() in extensions:
                        files.append(file_path)
                else:
                    files.append(file_path)
            
            processed_dirs += 1
        
        return files
    
    def get_git_history_cancellable(self, file_path: str = None, max_commits: int = 100, token: CancellationToken = None) -> List[Dict]:
        """
        Get git commit history with cancellation support
        
        Args:
            file_path: Specific file to get history for (optional)
            max_commits: Maximum number of commits to retrieve
            token: Cancellation token
            
        Returns:
            List of commit information dictionaries
        """
        if not self.repo:
            return []
        
        try:
            commits = []
            commit_iter = self.repo.iter_commits(paths=file_path, max_count=max_commits)
            
            for i, commit in enumerate(commit_iter):
                if token and i % 10 == 0:  # Check every 10 commits
                    token.check_cancellation()
                
                commits.append({
                    'hash': commit.hexsha,
                    'author': commit.author.name,
                    'email': commit.author.email,
                    'date': commit.committed_datetime,
                    'message': commit.message.strip(),
                    'files_changed': len(commit.stats.files)
                })
            
            return commits
        except OperationCancelledException:
            raise
        except Exception as e:
            # Completely suppress git history warnings for cloned repositories
            # Only log if it's a local repository that should have working git
            return []
    
    def add_rerun_button(self, analysis_type: str):
        """
        Previously added a rerun button, now removed as per requirements
        while maintaining functionality for cache clearing
        
        Args:
            analysis_type: Type of analysis to rerun
        """
        # Button removed but maintaining the functionality for cache management
        # This method is kept to maintain compatibility with existing code
        pass
    
    def get_analysis_with_control(self, analysis_type: str, loading_message: str):
        """
        Get analysis results with proper control flow - only runs when appropriate
        
        Args:
            analysis_type: Type of analysis
            loading_message: Message to show while loading
            
        Returns:
            Analysis results or None if analysis should not run
        """
        # Check if we should run analysis automatically or wait for user action
        analysis_running = st.session_state.get('analysis_running', False)
        
        # Check if this specific analyzer is running during parallel analysis
        analyzer_progress = st.session_state.get('analyzer_progress', {})
        is_analyzer_running = (analysis_type in analyzer_progress and 
                              analyzer_progress[analysis_type].get('status') == 'Running')
        
        # Check if analysis results are already available from parallel analysis
        if 'analysis_results' in st.session_state and analysis_type in st.session_state.analysis_results:
            result = st.session_state.analysis_results[analysis_type]
            if result.get('success', False) and 'analysis_data' in result:
                return result['analysis_data']
        
        # Only auto-analyze if analysis is explicitly running OR this analyzer is running in parallel
        if (analysis_running and st.session_state.get('analysis_started', False)) or is_analyzer_running:
            with self.display_loading_message(loading_message):
                return self.analyze()
        else:
            # Check for cached results first
            analysis = self.get_cached_analysis(analysis_type)
            
            if not analysis:
                # Show a button to run analysis manually
                st.info(f"Click 'Run Analysis' to {loading_message.lower()}.")
                if st.button("🚀 Run Analysis", key=f"run_{analysis_type}"):
                    with self.display_loading_message(loading_message):
                        return self.analyze()
                else:
                    return None  # Don't run analysis
            else:
                return analysis
    
    def display_parallel_ai_insights(self, analysis_type: str):
        """
        Display AI insights from parallel analysis if available
        
        Args:
            analysis_type: Type of analysis to show insights for
        """
        if 'parallel_ai_results' in st.session_state:
            results = st.session_state.parallel_ai_results
            
            # Map analysis types to result keys
            type_mapping = {
                'expertise_mapping': 'expertise',
                'timeline_analysis': 'timeline',
                'api_contracts': 'api_contracts',
                'ai_context': 'ai_context',
                'risk_analysis': 'risk_analysis',
                'development_patterns': 'development_patterns',
                'version_governance': 'version_governance',
                'tech_debt_detection': 'tech_debt',
                'design_patterns': 'design_patterns'
            }
            
            result_key = type_mapping.get(analysis_type, analysis_type)
            
            if result_key in results:
                result = results[result_key]
                
                st.subheader("🤖 AI Insights")
                
                if result.get('success', False) and result.get('insight'):
                    st.markdown("**AI-Generated Insights:**")
                    st.markdown(result['insight'])
                    
                    # Add timestamp
                    st.caption("Generated from parallel AI analysis")
                else:
                    st.error(f"AI analysis failed: {result.get('error', 'Unknown error')}")
                    
                return True  # Insights were displayed
        
        return False  # No insights available
    
    def display_loading_message(self, message: str):
        """Display a loading message"""
        return st.spinner(message)
    
    def display_error(self, message: str):
        """Display an error message"""
        st.error(message)
    
    def display_success(self, message: str):
        """Display a success message"""
        st.success(message)
    
    def display_info(self, message: str):
        """Display an info message"""
        st.info(message)
    
    def add_save_options(self, analysis_type: str, analysis_data: Dict[str, Any]):
        """
        Add save/export options for analysis results
        
        Args:
            analysis_type: Type of analysis
            analysis_data: Analysis data to save
        """
        st.subheader("💾 Save & Export Options")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Save as PDF
            pdf_key = f"pdf_prepared_{analysis_type}"
            
            if st.button("📋 Save as PDF", key=f"save_pdf_{analysis_type}"):
                try:
                    pdf_bytes = self._generate_single_analyzer_pdf(analysis_type, analysis_data)
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"{analysis_type}_report_{timestamp}.pdf"
                    
                    # Store PDF data in session state
                    st.session_state[pdf_key] = {
                        'data': pdf_bytes,
                        'filename': filename
                    }
                    st.rerun()
                except Exception as e:
                    st.error(f"PDF generation failed: {str(e)}")
            
            # Show download button if PDF is prepared
            if pdf_key in st.session_state:
                pdf_info = st.session_state[pdf_key]
                st.download_button(
                    label="⬇️ Download PDF Report",
                    data=pdf_info['data'],
                    file_name=pdf_info['filename'],
                    mime="application/pdf",
                    key=f"download_pdf_{analysis_type}",
                    on_click=lambda: st.session_state.pop(pdf_key, None)
                )
        
        with col2:
            # Save as DOCX
            docx_key = f"docx_prepared_{analysis_type}"
            
            if st.button("📄 Save as DOCX", key=f"save_docx_{analysis_type}"):
                try:
                    docx_bytes = self._generate_single_analyzer_docx(analysis_type, analysis_data)
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"{analysis_type}_report_{timestamp}.docx"
                    
                    # Store DOCX data in session state
                    st.session_state[docx_key] = {
                        'data': docx_bytes,
                        'filename': filename
                    }
                    st.rerun()
                except Exception as e:
                    st.error(f"DOCX generation failed: {str(e)}")
            
            # Show download button if DOCX is prepared
            if docx_key in st.session_state:
                docx_info = st.session_state[docx_key]
                st.download_button(
                    label="⬇️ Download DOCX Report",
                    data=docx_info['data'],
                    file_name=docx_info['filename'],
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_docx_{analysis_type}",
                    on_click=lambda: st.session_state.pop(docx_key, None)
                )
    
    
    def _generate_single_analyzer_pdf(self, analysis_type: str, analysis_data: Dict[str, Any]) -> bytes:
        """Generate a PDF report for a single analyzer"""
        import io
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
        
        # Create PDF buffer
        buffer = io.BytesIO()
        
        # Create the PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Define styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            textColor=colors.darkblue,
            alignment=TA_CENTER,
            spaceAfter=30
        )
        
        section_style = ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.darkblue,
            spaceBefore=20,
            spaceAfter=12
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            textColor=colors.black,
            alignment=TA_JUSTIFY,
            spaceAfter=8
        )
        
        # Build PDF content
        story = []
        
        # Title
        analyzer_title = analysis_type.replace('_', ' ').title()
        story.append(Paragraph(f"🔍 {analyzer_title} Report", title_style))
        story.append(Spacer(1, 20))
        
        # Repository information
        story.append(Paragraph("📁 Repository Information", section_style))
        repo_info_data = [
            ['Repository Path:', str(self.repo_path)],
            ['Analysis Date:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
            ['Analysis Type:', analyzer_title]
        ]
        
        repo_table = Table(repo_info_data, colWidths=[2.5*inch, 4*inch])
        repo_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightblue),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(repo_table)
        story.append(Spacer(1, 20))
        
        # Add AI insights if available
        if 'parallel_ai_results' in st.session_state:
            type_mapping = {
                'expertise_mapping': 'expertise',
                'timeline_analysis': 'timeline',
                'api_contracts': 'api_contracts',
                'ai_context': 'ai_context',
                'risk_analysis': 'risk_analysis',
                'development_patterns': 'development_patterns',
                'version_governance': 'version_governance',
                'tech_debt_detection': 'tech_debt',
                'design_patterns': 'design_patterns'
            }
            
            result_key = type_mapping.get(analysis_type, analysis_type)
            results = st.session_state.parallel_ai_results
            
            if result_key in results and results[result_key].get('success', False):
                story.append(Paragraph("🤖 AI-Generated Insights", section_style))
                
                insight_text = results[result_key].get('insight', '')
                if insight_text:
                    # Clean up the text for PDF
                    insight_text = insight_text.replace('\n', '<br/>')
                    insight_text = insight_text.replace('**', '<b>').replace('**', '</b>')
                    
                    # Split long text into paragraphs
                    paragraphs = insight_text.split('<br/><br/>')
                    for para in paragraphs:
                        if para.strip():
                            story.append(Paragraph(para.strip(), body_style))
                
                story.append(Spacer(1, 20))
        
        # Analysis results summary
        story.append(Paragraph("📊 Analysis Results", section_style))
        
        # Add key metrics
        key_metrics = []
        for key, value in analysis_data.items():
            if isinstance(value, (int, float)) or (isinstance(value, str) and len(str(value)) < 50):
                key_metrics.append([key.replace('_', ' ').title(), str(value)])
        
        if key_metrics and len(key_metrics) <= 15:  # Only show if reasonable number
            metrics_table = Table(key_metrics, colWidths=[2.5*inch, 3*inch])
            metrics_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            story.append(metrics_table)
        else:
            story.append(Paragraph("Detailed analysis data available in JSON export.", body_style))
        
        # Footer
        story.append(PageBreak())
        story.append(Paragraph(f"Report generated by AI-Powered Repository Analyzer - {analyzer_title} Module", 
                              ParagraphStyle('Footer', parent=styles['Normal'], 
                                           fontSize=10, textColor=colors.grey, 
                                           alignment=TA_CENTER)))
        
        # Build PDF
        doc.build(story)
        
        # Get PDF bytes
        buffer.seek(0)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        
        return pdf_bytes
    
    def _generate_single_analyzer_docx(self, analysis_type: str, analysis_data: Dict[str, Any]) -> bytes:
        """Generate a DOCX report for a single analyzer"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import RGBColor
            import io
            
            # Create a new document
            doc = Document()
            
            # Add title
            title = doc.add_heading(f"🔍 {analysis_type.replace('_', ' ').title()} Report", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add spacing
            doc.add_paragraph("")
            
            # Repository information section
            doc.add_heading("📁 Repository Information", level=1)
            
            repo_info = doc.add_paragraph()
            repo_info.add_run("Repository Path: ").bold = True
            repo_info.add_run(str(self.repo_path))
            
            date_info = doc.add_paragraph()
            date_info.add_run("Analysis Date: ").bold = True
            date_info.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
            
            type_info = doc.add_paragraph()
            type_info.add_run("Analysis Type: ").bold = True
            type_info.add_run(analysis_type.replace('_', ' ').title())
            
            doc.add_paragraph("")
            
            # Add AI insights if available
            if 'parallel_ai_results' in st.session_state:
                type_mapping = {
                    'expertise_mapping': 'expertise',
                    'timeline_analysis': 'timeline',
                    'api_contracts': 'api_contracts',
                    'ai_context': 'ai_context',
                    'risk_analysis': 'risk_analysis',
                    'development_patterns': 'development_patterns',
                    'version_governance': 'version_governance',
                    'tech_debt_detection': 'tech_debt',
                    'design_patterns': 'design_patterns'
                }
                
                result_key = type_mapping.get(analysis_type, analysis_type)
                results = st.session_state.parallel_ai_results
                
                if result_key in results and results[result_key].get('success', False):
                    doc.add_heading("🤖 AI-Generated Insights", level=1)
                    
                    insight_text = results[result_key].get('insight', '')
                    if insight_text:
                        # Clean up the text and split into paragraphs
                        paragraphs = insight_text.split('\n\n')
                        for para in paragraphs:
                            if para.strip():
                                # Handle basic markdown formatting
                                para = para.strip()
                                if para.startswith('**') and para.endswith('**'):
                                    # Bold paragraph
                                    p = doc.add_paragraph()
                                    p.add_run(para[2:-2]).bold = True
                                elif para.startswith('- '):
                                    # Bullet point
                                    doc.add_paragraph(para[2:], style='List Bullet')
                                else:
                                    # Regular paragraph
                                    doc.add_paragraph(para)
                    
                    doc.add_paragraph("")
            
            # Analysis results section
            doc.add_heading("📊 Analysis Results", level=1)
            
            # Add key metrics in a simple format
            self._add_data_to_docx(doc, analysis_data, level=2)
            
            # Add footer
            doc.add_page_break()
            footer_para = doc.add_paragraph(f"Report generated by AI-Powered Repository Analyzer - {analysis_type.replace('_', ' ').title()} Module")
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save to bytes
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            docx_bytes = docx_buffer.getvalue()
            docx_buffer.close()
            
            return docx_bytes
            
        except ImportError:
            raise Exception("python-docx package is required for DOCX generation. Please install it using: pip install python-docx")
        except Exception as e:
            raise Exception(f"DOCX generation failed: {str(e)}")
    
    def _add_data_to_docx(self, doc, data: Any, level: int = 2):
        """Add analysis data to DOCX document"""
        if isinstance(data, dict):
            for key, value in data.items():
                if key in ['error']:  # Skip error keys
                    continue
                
                # Add section heading - ensure key is string
                key_str = str(key)
                heading_text = key_str.replace('_', ' ').title()
                doc.add_heading(heading_text, level=level)
                
                if isinstance(value, dict):
                    self._add_data_to_docx(doc, value, level + 1)
                elif isinstance(value, list):
                    if value and isinstance(value[0], dict):
                        # Create a simple table for list of dictionaries
                        if len(value) > 0 and len(value) <= 10:  # Limit table size
                            headers = list(value[0].keys())[:5]  # Limit columns
                            
                            table = doc.add_table(rows=1, cols=len(headers))
                            table.style = 'Table Grid'
                            
                            # Add headers
                            header_cells = table.rows[0].cells
                            for i, header in enumerate(headers):
                                # Ensure header is string
                                header_str = str(header)
                                header_cells[i].text = header_str.replace('_', ' ').title()
                            
                            # Add data rows
                            for item in value[:10]:  # Limit to 10 rows
                                row_cells = table.add_row().cells
                                for i, header in enumerate(headers):
                                    cell_value = item.get(header, "")
                                    # Safely convert to string
                                    row_cells[i].text = str(cell_value)[:100]  # Limit cell content
                        else:
                            # Too much data, just show summary
                            doc.add_paragraph(f"Contains {len(value)} items (detailed data available in JSON export)")
                    else:
                        # Simple list
                        for item in value[:20]:  # Limit to 20 items
                            doc.add_paragraph(str(item)[:200], style='List Bullet')  # Limit item length
                        
                        if len(value) > 20:
                            doc.add_paragraph(f"... and {len(value) - 20} more items")
                else:
                    # Simple value
                    doc.add_paragraph(str(value)[:1000])  # Limit text length
                
                doc.add_paragraph("")  # Add spacing
        elif isinstance(data, list):
            for item in data[:10]:  # Limit to 10 items
                doc.add_paragraph(str(item)[:200], style='List Bullet')
            
            if len(data) > 10:
                doc.add_paragraph(f"... and {len(data) - 10} more items")
        else:
            doc.add_paragraph(str(data)[:1000])
