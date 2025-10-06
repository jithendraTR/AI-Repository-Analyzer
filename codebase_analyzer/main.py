"""
AI-Powered Repository Analyzer
Main application entry point with tabbed interface for different analysis types
"""

import streamlit as st
import os
from dotenv import load_dotenv, find_dotenv
import asyncio
import concurrent.futures
from typing import Dict, Any
import threading
import time
import json
import io
from datetime import datetime

# Document generation imports
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown

# Load environment variables from parent directory
load_dotenv(find_dotenv())

# Add parent directory to path for imports
import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Import analysis modules
from analyzers.commit_mapping import ExpertiseMapper
from analyzers.timeline_analysis import TimelineAnalyzer
from analyzers.api_analysis import APIContractAnalyzer
from analyzers.ai_integration_analysis import AIContextAnalyzer
from analyzers.risk_analysis import RiskAnalysisAnalyzer
from utils.ai_client import OpenArenaClient

class ParallelAIAnalyzer:
    """Handles parallel AI analysis for all tabs with cancellation support"""
    
    def __init__(self, repo_path: str):
        self.repo_path = repo_path
        self.ai_client = OpenArenaClient()
        self.analyzers = {
            'expertise': ExpertiseMapper(repo_path),
            'timeline': TimelineAnalyzer(repo_path),
            'api_contracts': APIContractAnalyzer(repo_path),
            'ai_context': AIContextAnalyzer(repo_path),
            'risk_analysis': RiskAnalysisAnalyzer(repo_path)
        }
        self.cancellation_token = None
    
    def set_cancellation_token(self, token):
        """Set the cancellation token for this analyzer"""
        self.cancellation_token = token
    
    def run_single_analysis(self, analyzer_name: str, analyzer) -> Dict[str, Any]:
        """Run analysis and generate AI insight for a specific analyzer with cancellation support"""
        try:
            # Check for cancellation before starting
            if self.cancellation_token and self.cancellation_token.is_cancelled():
                return {
                    "analyzer": analyzer_name,
                    "error": "Operation was cancelled",
                    "success": False,
                    "cancelled": True
                }
            
            # Progress callback for individual analyzer steps
            def analyzer_progress(step, total, status):
                print(f"DEBUG: {analyzer_name} - Step {step}/{total}: {status}")
            
            # Get analysis data - this is now running in parallel with progress tracking
            analysis_data = analyzer.analyze(token=self.cancellation_token, progress_callback=analyzer_progress)
            
            # Check for cancellation after analysis
            if self.cancellation_token and self.cancellation_token.is_cancelled():
                return {
                    "analyzer": analyzer_name,
                    "error": "Operation was cancelled",
                    "success": False,
                    "cancelled": True
                }
            
            if "error" in analysis_data:
                return {"analyzer": analyzer_name, "error": analysis_data["error"]}
            
            # Generate appropriate prompt based on analyzer type
            prompt = self._generate_prompt(analyzer_name, analysis_data)
            
            # Check for cancellation before AI query
            if self.cancellation_token and self.cancellation_token.is_cancelled():
                return {
                    "analyzer": analyzer_name,
                    "error": "Operation was cancelled",
                    "success": False,
                    "cancelled": True
                }
            
            # Get AI insight
            insight = self.ai_client.query(prompt)
            
            return {
                "analyzer": analyzer_name,
                "insight": insight,
                "success": True,
                "analysis_data": analysis_data  # Include raw analysis data
            }
        except Exception as e:
            # Check if it's a cancellation
            if "cancelled" in str(e).lower():
                return {
                    "analyzer": analyzer_name,
                    "error": "Operation was cancelled",
                    "success": False,
                    "cancelled": True
                }
            return {
                "analyzer": analyzer_name,
                "error": str(e),
                "success": False
            }
    
    def _generate_prompt(self, analyzer_name: str, analysis_data: Dict[str, Any]) -> str:
        """Generate appropriate prompt for each analyzer type"""
        
        prompts = {
            'expertise': f"""
            Analyze this expertise mapping data and provide insights:
            {str(analysis_data)[:2000]}
            
            Provide:
            1. Key expertise areas and knowledge gaps
            2. Team collaboration patterns
            3. Knowledge transfer recommendations
            4. Risk assessment for key person dependencies
            """,
            
            'timeline': f"""
            Analyze this project timeline data:
            Total commits: {analysis_data.get('total_commits', 0)}
            Recent activity: {analysis_data.get('recent_changes', {}).get('total_recent_commits', 0)}
            
            Provide:
            1. Project maturity assessment
            2. Development velocity trends
            3. Team activity patterns
            4. Timeline management recommendations
            """,
            
            'api_contracts': f"""
            Analyze these API contracts and integration points:
            {str(analysis_data)[:2000]}
            
            Provide:
            1. API design quality assessment
            2. Integration complexity analysis
            3. Potential breaking changes risks
            4. API evolution recommendations
            """,
            
            'ai_context': f"""
            Analyze this codebase context for AI feature integration:
            {str(analysis_data)[:2000]}
            
            Provide:
            1. Best locations for new AI features
            2. Integration complexity assessment
            3. Architectural recommendations
            4. Implementation strategy
            """,
            
            'risk_analysis': f"""
            Analyze these risk factors:
            {str(analysis_data)[:2000]}
            
            Provide:
            1. Critical risk assessment
            2. Test coverage recommendations
            3. Security vulnerability priorities
            4. Risk mitigation strategies
            """,
            
            'development_patterns': f"""
            Analyze these development patterns:
            {str(analysis_data)[:2000]}
            
            Provide:
            1. Code quality assessment
            2. Pattern consistency analysis
            3. Best practices recommendations
            4. Refactoring priorities
            """,
            
            'version_governance': f"""
            Analyze version and dependency management:
            {str(analysis_data)[:2000]}
            
            Provide:
            1. Dependency health assessment
            2. Version conflict risks
            3. Update strategy recommendations
            4. Security vulnerability priorities
            """,
            
            'tech_debt': f"""
            Analyze technical debt indicators:
            {str(analysis_data)[:2000]}
            
            Provide:
            1. Technical debt severity assessment
            2. Refactoring priorities
            3. Code quality improvement plan
            4. Long-term maintenance strategy
            """,
            
            'design_patterns': f"""
            Analyze design pattern usage and deviations:
            {str(analysis_data)[:2000]}
            
            Provide:
            1. Design pattern compliance assessment
            2. Architecture consistency analysis
            3. Pattern improvement recommendations
            4. Code structure optimization suggestions
            """
        }
        
        return prompts.get(analyzer_name, f"Analyze this data: {str(analysis_data)[:2000]}")
    
    def run_parallel_analysis(self, progress_callback=None) -> Dict[str, Any]:
        """Run AI analysis for all analyzers in parallel with cancellation support"""
        results = {}
        completed_count = 0
        total_count = len(self.analyzers)
        
        # Debug: Print what analyzers we're about to run
        print(f"DEBUG: Starting parallel analysis for {total_count} analyzers: {list(self.analyzers.keys())}")
        
        # Use a smaller thread pool to avoid overwhelming the system
        max_workers = min(total_count, 2)  # Reduced to 2 to prevent timeouts
        print(f"DEBUG: Using {max_workers} worker threads")
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all tasks at once - this should start them all simultaneously
            future_to_analyzer = {}
            
            # Submit all tasks and track them
            for name, analyzer in self.analyzers.items():
                print(f"DEBUG: Submitting {name} to thread pool")
                future = executor.submit(self.run_single_analysis, name, analyzer)
                future_to_analyzer[future] = name
            
            print(f"DEBUG: All {len(future_to_analyzer)} tasks submitted to thread pool")
            
            # Update progress to show all analyzers are starting
            if progress_callback:
                analyzer_names = [name.replace('_', ' ').title() for name in self.analyzers.keys()]
                progress_callback(0, total_count, f"Starting: {', '.join(analyzer_names[:3])}{'...' if len(analyzer_names) > 3 else ''}")
            
            # Collect results as they complete with timeout
            try:
                for future in concurrent.futures.as_completed(future_to_analyzer, timeout=900):  # 15 minute total timeout
                    # Check for cancellation
                    if self.cancellation_token and self.cancellation_token.is_cancelled():
                        print("DEBUG: Cancellation detected, stopping remaining tasks")
                        # Cancel remaining futures
                        for remaining_future in future_to_analyzer:
                            if not remaining_future.done():
                                remaining_future.cancel()
                        
                        # Add cancelled status to remaining analyzers
                        for remaining_future, remaining_analyzer in future_to_analyzer.items():
                            if not remaining_future.done():
                                results[remaining_analyzer] = {
                                    "analyzer": remaining_analyzer,
                                    "error": "Operation was cancelled",
                                    "success": False,
                                    "cancelled": True
                                }
                        break
                    
                    analyzer_name = future_to_analyzer[future]
                    print(f"DEBUG: {analyzer_name} completed")
                    
                    try:
                        result = future.result(timeout=5)  # Short timeout since future is already done
                        results[analyzer_name] = result
                        completed_count += 1
                        
                        # Update progress if callback provided
                        if progress_callback:
                            # Show which analyzers are still running
                            remaining_analyzers = [name for f, name in future_to_analyzer.items() if not f.done()]
                            if remaining_analyzers:
                                progress_callback(completed_count, total_count, f"Running: {', '.join(remaining_analyzers[:2])}{'...' if len(remaining_analyzers) > 2 else ''}")
                            else:
                                progress_callback(completed_count, total_count, "Finalizing results")
                            
                    except concurrent.futures.TimeoutError:
                        print(f"DEBUG: {analyzer_name} result timed out")
                        results[analyzer_name] = {
                            "analyzer": analyzer_name,
                            "error": "Result retrieval timed out",
                            "success": False
                        }
                        completed_count += 1
                    except Exception as e:
                        print(f"DEBUG: {analyzer_name} failed with error: {str(e)}")
                        results[analyzer_name] = {
                            "analyzer": analyzer_name,
                            "error": str(e),
                            "success": False
                        }
                        completed_count += 1
                        
            except concurrent.futures.TimeoutError:
                print("DEBUG: Overall analysis timed out")
                # Handle any remaining futures that didn't complete
                for future, analyzer_name in future_to_analyzer.items():
                    if analyzer_name not in results:
                        results[analyzer_name] = {
                            "analyzer": analyzer_name,
                            "error": "Analysis timed out",
                            "success": False
                        }
        
        print(f"DEBUG: Parallel analysis completed. Results: {len(results)} analyzers")
        return results

def generate_pdf_report(analysis_results: Dict[str, Any], repo_path: str) -> io.BytesIO:
    """Generate a PDF report from analysis results"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    
    # Get sample style sheet
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=1  # Center alignment
    )
    story.append(Paragraph("AI Codebase Analysis Report", title_style))
    story.append(Spacer(1, 12))
    
    # Repository info
    info_style = styles['Normal']
    story.append(Paragraph(f"<b>Repository:</b> {repo_path}", info_style))
    story.append(Paragraph(f"<b>Generated:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", info_style))
    story.append(Spacer(1, 20))
    
    # Analysis sections
    section_style = ParagraphStyle(
        'SectionTitle',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=12,
        textColor=colors.darkblue
    )
    
    analysis_names = {
        'expertise': 'Team Expertise & Commit Mapping',
        'timeline': 'Timeline Analysis',
        'api_contracts': 'API Analysis',
        'ai_context': 'AI Integration Analysis',
        'risk_analysis': 'Risk Analysis'
    }
    
    for analyzer_key, result in analysis_results.items():
        if result.get('success', False):
            # Section title
            section_name = analysis_names.get(analyzer_key, analyzer_key.replace('_', ' ').title())
            story.append(Paragraph(section_name, section_style))
            
            # AI Insights
            insight_text = result.get('insight', 'No insights available')
            # Clean up text for PDF
            insight_text = insight_text.replace('**', '<b>').replace('**', '</b>')
            insight_text = insight_text.replace('*', '<i>').replace('*', '</i>')
            
            # Split into paragraphs
            paragraphs = insight_text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.strip(), info_style))
                    story.append(Spacer(1, 6))
            
            story.append(Spacer(1, 12))
        else:
            # Error section
            section_name = analysis_names.get(analyzer_key, analyzer_key.replace('_', ' ').title())
            story.append(Paragraph(f"{section_name} - Error", section_style))
            error_text = result.get('error', 'Unknown error occurred')
            story.append(Paragraph(f"Error: {error_text}", info_style))
            story.append(Spacer(1, 12))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer

def generate_docx_report(analysis_results: Dict[str, Any], repo_path: str) -> io.BytesIO:
    """Generate a DOCX report from analysis results"""
    doc = Document()
    
    # Title
    title = doc.add_heading('AI Codebase Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Repository info
    doc.add_paragraph(f"Repository: {repo_path}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()  # Empty paragraph for spacing
    
    analysis_names = {
        'expertise': 'Team Expertise & Commit Mapping',
        'timeline': 'Timeline Analysis', 
        'api_contracts': 'API Analysis',
        'ai_context': 'AI Integration Analysis',
        'risk_analysis': 'Risk Analysis'
    }
    
    for analyzer_key, result in analysis_results.items():
        if result.get('success', False):
            # Section title
            section_name = analysis_names.get(analyzer_key, analyzer_key.replace('_', ' ').title())
            doc.add_heading(section_name, level=1)
            
            # AI Insights
            insight_text = result.get('insight', 'No insights available')
            
            # Split into paragraphs and add to document
            paragraphs = insight_text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    # Handle basic markdown formatting
                    para_text = para.strip()
                    if para_text.startswith('**') and para_text.endswith('**'):
                        # Bold paragraph
                        p = doc.add_paragraph()
                        run = p.add_run(para_text[2:-2])
                        run.bold = True
                    elif para_text.startswith('*') and para_text.endswith('*'):
                        # Italic paragraph
                        p = doc.add_paragraph()
                        run = p.add_run(para_text[1:-1])
                        run.italic = True
                    else:
                        doc.add_paragraph(para_text)
            
            doc.add_paragraph()  # Add spacing after section
        else:
            # Error section
            section_name = analysis_names.get(analyzer_key, analyzer_key.replace('_', ' ').title())
            doc.add_heading(f"{section_name} - Error", level=1)
            error_text = result.get('error', 'Unknown error occurred')
            doc.add_paragraph(f"Error: {error_text}")
            doc.add_paragraph()  # Add spacing after section
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.set_page_config(
        page_title="AI Codebase Analyzer",
        page_icon="🔍",
        layout="wide"
    )
    
    # Initialize sidebar state
    if 'sidebar_collapsed' not in st.session_state:
        st.session_state.sidebar_collapsed = False
    
    # Hide Streamlit's default deploy button and menu
    hide_streamlit_style = f"""
    <style>
    #MainMenu {{visibility: hidden;}}
    .stDeployButton {{display:none !important;}}
    .stActionButton {{display:none !important;}}
    [data-testid="stToolbar"] {{display: none !important;}}
    [data-testid="stDecoration"] {{display: none !important;}}
    [data-testid="stStatusWidget"] {{display: none !important;}}
    .stApp > header {{display: none !important;}}
    .stApp > .main .block-container {{padding-top: 2rem !important;}}
    footer {{visibility: hidden !important;}}
    #stDecoration {{display:none !important;}}
    button[title="Deploy this app"] {{display: none !important;}}
    .css-1rs6os {{display: none !important;}}
    .css-17eq0hr {{display: none !important;}}
    
    /* Hide Streamlit's default sidebar collapse button and other unwanted elements */
    [data-testid="stSidebarCollapseButton"] {{display: none !important;}}
    [data-testid="stBaseButton-headerNoPadding"] {{display: none !important;}}
    button[kind="headerNoPadding"] {{display: none !important;}}
    [data-testid="stLogoSpacer"] {{display: none !important;}}
    
    /* Hide sidebar when collapsed */
    {'section[data-testid="stSidebar"] { display: none !important; }' if st.session_state.sidebar_collapsed else ''}
    
    /* Adjust main content when sidebar is collapsed */
    {'div[data-testid="stAppViewContainer"] .main .block-container { margin-left: 0 !important; max-width: none !important; }' if st.session_state.sidebar_collapsed else ''}
    
    /* Custom toggle button styling */
    .sidebar-toggle {{
        position: fixed;
        top: 1rem;
        left: {'0.5rem' if st.session_state.sidebar_collapsed else '21rem'};
        z-index: 999999;
        transition: left 0.3s ease;
    }}
    
    .sidebar-toggle .stButton > button {{
        background: #262730 !important;
        color: white !important;
        border: 1px solid #464853 !important;
        border-radius: 0.25rem !important;
        padding: 0.25rem 0.5rem !important;
        font-size: 16px !important;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1) !important;
        transition: all 0.3s ease !important;
        min-height: auto !important;
        height: 2rem !important;
    }}
    
    .sidebar-toggle .stButton > button:hover {{
        background: #464853 !important;
        box-shadow: 0 4px 8px rgba(0,0,0,0.2) !important;
        transform: translateY(-1px) !important;
    }}
    
    .sidebar-toggle .stButton > button:active {{
        transform: translateY(0px) !important;
        box-shadow: 0 1px 2px rgba(0,0,0,0.1) !important;
    }}
    </style>
    """
    st.markdown(hide_streamlit_style, unsafe_allow_html=True)
    
    # Create sidebar toggle button when sidebar is collapsed (floating)
    if st.session_state.sidebar_collapsed:
        toggle_container = st.container()
        with toggle_container:
            st.markdown('<div class="sidebar-toggle">', unsafe_allow_html=True)
            if st.button("►", key="open_sidebar", help="Open Sidebar"):
                st.session_state.sidebar_collapsed = False
                st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)
    
    st.title("🔍 AI-Powered Codebase Analyzer")
    st.markdown("Accelerate codebase onboarding and architectural discovery")
    
    # Sidebar for repository selection
    with st.sidebar:
        # Add close sidebar button at the top
        col1, col2 = st.columns([3, 1])
        with col1:
            st.header("Repository Configuration")
        with col2:
            if st.button("◄", key="close_sidebar_internal", help="Close Sidebar", use_container_width=False):
                st.session_state.sidebar_collapsed = True
                st.rerun()
        
        # Add custom CSS for spacing around text input and styling the placeholder text
        st.markdown("""
        <style>
        .stTextInput > div > div > input {
            margin-bottom: 10px !important;
        }
        .stTextInput > label {
            margin-bottom: 5px !important;
            color: black !important; /* Set label to black */
        }
        .stTextInput {
            margin-bottom: 15px !important;
        }
        
        /* Ensure label text stays black */
        .stTextInput label {
            color: black !important;
        }
        
        /* ONLY target "Press Enter to apply" text - light green color */
        /* Be very specific to avoid affecting the label */
        .stTextInput small {
            color: #90EE90 !important;
        }
        div[data-testid="stTextInput"] small {
            color: #90EE90 !important;
        }
        
        /* Target helper text that appears below input */
        .stTextInput div[data-baseweb="input"] + div small {
            color: #90EE90 !important;
        }
        .stTextInput div[data-baseweb="input"] ~ div small {
            color: #90EE90 !important;
        }
        
        /* Target any small text elements that might contain helper text */
        .stTextInput div small:contains("Press Enter") {
            color: #90EE90 !important;
        }
        .stTextInput div small:contains("apply") {
            color: #90EE90 !important;
        }
        
        /* Specific targeting for input helper text positioned below input field */
        .stTextInput > div > div[data-baseweb="input"] + * small {
            color: #90EE90 !important;
        }
        
        /* Target any text that might be helper text but avoid affecting labels */
        .stTextInput div:not(label) small {
            color: #90EE90 !important;
        }
        
        /* Additional safety - ensure label remains black */
        .stTextInput > label,
        .stTextInput label {
            color: black !important;
        }
        </style>
        """, unsafe_allow_html=True)
        
        repo_path = st.text_input("Repository Path", placeholder="/path/to/your/repo")
        
        # Add line gap between input field and validation text
        st.markdown("")
        st.markdown("")
        
        # Auto-validate repository path without button
        if repo_path and os.path.exists(repo_path):
            if 'repo_path' not in st.session_state or st.session_state.repo_path != repo_path:
                st.session_state.repo_path = repo_path
                st.success("Repository loaded successfully!")
        elif repo_path and not os.path.exists(repo_path):
            st.error("Please provide a valid repository path")
        
        # Analysis Selection Section - Always visible when repo path is provided
        if repo_path and os.path.exists(repo_path):
            st.header("🎯 Select Analysis")
            st.markdown("Choose which analyses to run on your repository")
            
            analysis_options = {
                'expertise': '👥 Commits Mapping',
                'timeline': '📅 Timeline Analysis',
                'api_contracts': '🔌 API Analysis',
                'ai_context': '🤖 AI Integration Analysis',
                'risk_analysis': '⚠️ Risk Analysis'
            }
            
            # Initialize session state for selected analyses if not exists
            if 'selected_analyses' not in st.session_state:
                st.session_state.selected_analyses = []
            
            # Quick selection buttons
            col1, col2 = st.columns(2)
            with col1:
                if st.button("📊 Select All", key="select_all"):
                    st.session_state.selected_analyses = list(analysis_options.keys())
                    st.rerun()
            with col2:
                if st.button("🗑️ Clear All", key="clear_all"):
                    st.session_state.selected_analyses = []
                    st.rerun()
            
            # Use session state as the value for multiselect
            selected_analyses = st.multiselect(
                "Select analyses to run:",
                options=list(analysis_options.keys()),
                format_func=lambda x: analysis_options[x],
                default=st.session_state.selected_analyses,
                help="Choose one or more analyses to run on your repository"
            )
            
            # Update session state when multiselect changes
            if st.session_state.selected_analyses != selected_analyses:
                st.session_state.selection_changing = True
                st.session_state.selected_analyses = selected_analyses
                # Clear the flag after a brief moment
                import threading
                def clear_flag():
                    import time
                    time.sleep(0.1)
                    if 'selection_changing' in st.session_state:
                        del st.session_state.selection_changing
                threading.Thread(target=clear_flag, daemon=True).start()
            else:
                # Ensure flag is cleared if selections are the same
                if 'selection_changing' in st.session_state:
                    del st.session_state.selection_changing
            
            # Run Analysis Section
            if selected_analyses:
                st.markdown("---")
                st.subheader("🚀 Run Analysis")
                
                    # Check if analysis is running
                analysis_running = st.session_state.get('analysis_running', False)
                
                if not analysis_running:
                    if st.button(f"🚀 Run {len(selected_analyses)} Selected Analysis{'es' if len(selected_analyses) > 1 else ''}", key="run_selected"):
                        # Mark as running and explicitly set button clicked flag
                        st.session_state.analysis_running = True
                        st.session_state.analysis_button_clicked = True
                        st.session_state.analysis_token = f"analysis_{int(time.time())}"
                        st.rerun()
                else:
                    # Show progress and stop button
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    # Stop button aligned properly
                    if st.button("🛑 Stop Analysis", key="stop_analysis", use_container_width=True):
                        st.session_state.analysis_cancelled = True
                        st.session_state.analysis_running = False
                        # Clear any existing results to avoid confusion
                        if 'analysis_results' in st.session_state:
                            del st.session_state.analysis_results
                        st.warning("Analysis stopped. You can start a new analysis now.")
                        st.rerun()
                    
                    # Check if analysis was cancelled immediately
                    if st.session_state.get('analysis_cancelled', False):
                        # Analysis was cancelled - clean up immediately
                        st.session_state.analysis_running = False
                        cleanup_keys = ['analysis_cancelled', 'analysis_token', 'analysis_button_clicked']
                        for key in cleanup_keys:
                            if key in st.session_state:
                                del st.session_state[key]
                        # Clear any partial results
                        if 'analysis_results' in st.session_state:
                            del st.session_state.analysis_results
                        st.info("Analysis was stopped. You can start a new analysis now.")
                        st.rerun()
                    else:
                        # Initialize analysis if not started AND button was clicked
                        if 'analysis_started' not in st.session_state and st.session_state.get('analysis_button_clicked', False):
                            st.session_state.analysis_started = True
                            st.session_state.analysis_progress = 0
                            st.session_state.analysis_status = "Starting parallel analysis..."
                            st.session_state.analysis_start_time = time.time()
                            
                            # Initialize progress tracking for each analyzer
                            st.session_state.analyzer_progress = {}
                            for analyzer_name in selected_analyses:
                                st.session_state.analyzer_progress[analyzer_name] = {
                                    'status': 'Queued',
                                    'progress': 0,
                                    'step': 0,
                                    'total_steps': 0
                                }
                            
                            # Run analysis directly without threading to avoid UI update issues
                            try:
                                parallel_analyzer = ParallelAIAnalyzer(st.session_state.repo_path)
                                
                                # Create cancellation token
                                from analyzers.base_analyzer import CancellationToken
                                token = CancellationToken(st.session_state.analysis_token)
                                parallel_analyzer.set_cancellation_token(token)
                                
                                # Filter analyzers to only selected ones
                                filtered_analyzers = {
                                    key: analyzer for key, analyzer in parallel_analyzer.analyzers.items()
                                    if key in selected_analyses
                                }
                                parallel_analyzer.analyzers = filtered_analyzers
                                
                                # Progress callback that updates session state
                                def update_progress(completed, total, current_analyzer):
                                    if not st.session_state.get('analysis_cancelled', False):
                                        progress = completed / total if total > 0 else 0
                                        st.session_state.analysis_progress = progress
                                        
                                        if completed == 0:
                                            st.session_state.analysis_status = f"Starting analyzers: {', '.join([name.replace('_', ' ').title() for name in selected_analyses[:2]])}{'...' if len(selected_analyses) > 2 else ''}"
                                        elif completed < total:
                                            st.session_state.analysis_status = f"Running: {current_analyzer.replace('_', ' ').title()}... ({completed}/{total})"
                                        else:
                                            st.session_state.analysis_status = f"Finalizing results... ({completed}/{total})"
                                        
                                        # Update individual analyzer status
                                        for analyzer_name in selected_analyses:
                                            if analyzer_name in st.session_state.analyzer_progress:
                                                if current_analyzer == analyzer_name:
                                                    st.session_state.analyzer_progress[analyzer_name]['status'] = 'Running'
                                                elif completed > 0:
                                                    # Check if this analyzer might be completed (rough estimation)
                                                    analyzer_index = list(selected_analyses).index(analyzer_name) if analyzer_name in selected_analyses else -1
                                                    if analyzer_index >= 0 and analyzer_index < completed:
                                                        st.session_state.analyzer_progress[analyzer_name]['status'] = 'Completed'
                                                        st.session_state.analyzer_progress[analyzer_name]['progress'] = 100
                                
                                # Run parallel analysis - this will execute multiple analyzers simultaneously
                                results = parallel_analyzer.run_parallel_analysis(progress_callback=update_progress)
                                
                                # Store results if not cancelled
                                if not st.session_state.get('analysis_cancelled', False):
                                    st.session_state.analysis_results = results
                                    st.session_state.analysis_completed = True
                                    st.session_state.analysis_progress = 1.0
                                    end_time = time.time()
                                    st.session_state.analysis_status = f"Completed in {end_time - st.session_state.analysis_start_time:.1f} seconds"
                                    
                                    # Update individual analyzer statuses
                                    for analyzer_name, result in results.items():
                                        if analyzer_name in st.session_state.analyzer_progress:
                                            if result.get('success', False):
                                                st.session_state.analyzer_progress[analyzer_name]['status'] = 'Completed'
                                                st.session_state.analyzer_progress[analyzer_name]['progress'] = 100
                                            else:
                                                st.session_state.analyzer_progress[analyzer_name]['status'] = 'Failed'
                            
                            except Exception as e:
                                if not st.session_state.get('analysis_cancelled', False):
                                    st.session_state.analysis_error = str(e)
                                    st.session_state.analysis_completed = True
                        
                        # Update UI with current progress
                        current_progress = st.session_state.get('analysis_progress', 0)
                        current_status = st.session_state.get('analysis_status', "Starting...")
                        
                        progress_bar.progress(current_progress)
                        status_text.text(current_status)
                        
                        
                        # Force UI refresh during analysis
                        if not st.session_state.get('analysis_completed', False):
                            time.sleep(0.5)  # Brief pause to allow UI updates
                            st.rerun()
                        
                        # Check if analysis is completed
                        if st.session_state.get('analysis_completed', False):
                            # Analysis finished
                            if 'analysis_results' in st.session_state:
                                results = st.session_state.analysis_results
                                successful = sum(1 for r in results.values() if r.get('success', False))
                                cancelled = sum(1 for r in results.values() if r.get('cancelled', False))
                                total = len(results)
                                
                                if cancelled > 0:
                                    st.warning(f"⚠️ Analysis stopped: {successful}/{total} completed, {cancelled} cancelled")
                                elif successful == total:
                                    st.success(f"🎉 All {total} analyses completed successfully!")
                                else:
                                    st.warning(f"⚠️ {successful}/{total} analyses completed successfully")
                            elif 'analysis_error' in st.session_state:
                                st.error(f"Analysis failed: {st.session_state.analysis_error}")
                            
                            # Clean up
                            st.session_state.analysis_running = False
                            cleanup_keys = ['analysis_started', 'analysis_progress', 'analysis_status', 
                                          'analysis_completed', 'analysis_start_time', 'analysis_thread', 'analysis_error',
                                          'analysis_button_clicked', 'analyzer_progress']
                            for key in cleanup_keys:
                                if key in st.session_state:
                                    del st.session_state[key]
                            if 'analysis_cancelled' in st.session_state:
                                del st.session_state.analysis_cancelled
                            if 'analysis_token' in st.session_state:
                                del st.session_state.analysis_token
                            
                            time.sleep(1)  # Brief pause to show results
                            st.rerun()
                        else:
                            # Analysis still running, refresh every 1 second
                            # Only rerun if analysis is actually running to avoid auto-start
                            if st.session_state.get('analysis_started', False):
                                time.sleep(1)
                                st.rerun()
            
            # Show analysis results summary
            if 'analysis_results' in st.session_state:
                st.markdown("---")
                st.subheader("📊 Analysis Results")
                results = st.session_state.analysis_results
                successful = sum(1 for r in results.values() if r.get('success', False))
                total = len(results)
                st.success(f"✅ Analysis Complete: {successful}/{total} successful")
                
                # Show brief status for each analyzer
                for analyzer_name, result in results.items():
                    if result.get('success', False):
                        st.success(f"✅ {analyzer_name.replace('_', ' ').title()}")
                    else:
                        st.error(f"❌ {analyzer_name.replace('_', ' ').title()}: {result.get('error', 'Unknown error')}")
                
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("🗑️ Clear Results", key="clear_results"):
                        # Store results for potential restoration before clearing
                        st.session_state.last_cleared_results = st.session_state.analysis_results.copy()
                        st.session_state.results_cleared_timestamp = time.time()
                        del st.session_state.analysis_results
                        st.success("Results cleared! Use 'Reopen Results' to restore.")
                        st.rerun()
                with col2:
                    if st.button("◄ Close Sidebar", key="collapse_sidebar_from_results"):
                        st.session_state.sidebar_collapsed = True
                        st.rerun()
            
            # Show reopen button if results were recently cleared
            elif 'last_cleared_results' in st.session_state:
                st.markdown("---")
                st.subheader("🔄 Restore Results")
                
                # Show when results were cleared
                if 'results_cleared_timestamp' in st.session_state:
                    cleared_time = time.time() - st.session_state.results_cleared_timestamp
                    if cleared_time < 60:
                        st.info(f"Results cleared {int(cleared_time)} seconds ago")
                    else:
                        st.info(f"Results cleared {int(cleared_time/60)} minutes ago")
                
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("🔄 Reopen Results", key="reopen_results"):
                        # Restore the cleared results
                        st.session_state.analysis_results = st.session_state.last_cleared_results.copy()
                        del st.session_state.last_cleared_results
                        if 'results_cleared_timestamp' in st.session_state:
                            del st.session_state.results_cleared_timestamp
                        st.success("Results restored!")
                        st.rerun()
                with col2:
                    if st.button("🗑️ Permanently Delete", key="permanent_delete"):
                        del st.session_state.last_cleared_results
                        if 'results_cleared_timestamp' in st.session_state:
                            del st.session_state.results_cleared_timestamp
                        st.success("Results permanently deleted.")
                        st.rerun()
            
        
        # Legacy AI Analysis Section (kept for backward compatibility but hidden)
        if False and 'repo_path' in st.session_state:
            st.header("🤖 AI Analysis")
            st.markdown("Generate AI insights for all tabs simultaneously")
            
            # Show status of parallel analysis
            if 'parallel_ai_results' in st.session_state:
                results = st.session_state.parallel_ai_results
                successful = sum(1 for r in results.values() if r.get('success', False))
                total = len(results)
                st.success(f"✅ AI Analysis Complete: {successful}/{total} successful")
                
                # Show brief status for each analyzer
                for analyzer_name, result in results.items():
                    if result.get('success', False):
                        st.success(f"✅ {analyzer_name.replace('_', ' ').title()}")
                    else:
                        st.error(f"❌ {analyzer_name.replace('_', ' ').title()}: {result.get('error', 'Unknown error')}")
            
            # Selective AI analysis
            st.subheader("🎯 Selective Analysis")
            analysis_options = {
                'expertise': '👥 Commits Mapping',
                'timeline': '📅 Timeline Analysis',
                'api_contracts': '🔌 API Analysis',
                'ai_context': '🤖 AI Integration Analysis',
                'risk_analysis': '⚠️ Risk Analysis',
                'development_patterns': '🏗️ Development Patterns',
                'version_governance': '📦 Version Governance',
                'tech_debt': '🔧 Tech Debt Detection',
                'design_patterns': '📐 Design Patterns'
            }
            
            selected_analyses = st.multiselect(
                "Select analyses to run:",
                options=list(analysis_options.keys()),
                format_func=lambda x: analysis_options[x],
                default=list(analysis_options.keys())
            )
            
            # Check if selective analysis is running
            selective_running = st.session_state.get('selective_analysis_running', False)
            
            if not selective_running:
                if st.button("🚀 Run Selected Analyses", key="selective_ai"):
                    if selected_analyses:
                        # Mark as running
                        st.session_state.selective_analysis_running = True
                        st.session_state.selective_analysis_token = f"selective_{int(time.time())}"
                        st.rerun()
                    else:
                        st.warning("Please select at least one analysis to run")
            else:
                # Show progress and stop button (no columns in sidebar)
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                if st.button("🛑 Stop Selected Analysis", key="stop_selective"):
                    st.session_state.selective_analysis_cancelled = True
                    st.warning("Stopping analysis...")
                    st.rerun()
                
                # Run the analysis
                if not st.session_state.get('selective_analysis_cancelled', False):
                    parallel_analyzer = ParallelAIAnalyzer(st.session_state.repo_path)
                    
                    # Create cancellation token
                    from analyzers.base_analyzer import CancellationToken
                    token = CancellationToken(st.session_state.selective_analysis_token)
                    parallel_analyzer.set_cancellation_token(token)
                    
                    # Filter analyzers to only selected ones
                    filtered_analyzers = {
                        key: analyzer for key, analyzer in parallel_analyzer.analyzers.items()
                        if key in selected_analyses
                    }
                    parallel_analyzer.analyzers = filtered_analyzers
                    
                    # Progress callback
                    def update_progress(completed, total, current_analyzer):
                        progress = (completed / total) * 100
                        progress_bar.progress(progress / 100.0)
                        status_text.text(f"Analyzing {current_analyzer.replace('_', ' ').title()}... ({completed}/{total})")
                    
                    # Run selective analysis
                    start_time = time.time()
                    results = parallel_analyzer.run_parallel_analysis(progress_callback=update_progress)
                    end_time = time.time()
                    
                    # Update progress
                    progress_bar.progress(1.0)
                    status_text.text(f"Completed in {end_time - start_time:.1f} seconds")
                    
                    # Merge with existing results
                    if 'parallel_ai_results' not in st.session_state:
                        st.session_state.parallel_ai_results = {}
                    st.session_state.parallel_ai_results.update(results)
                    
                    # Show summary
                    successful = sum(1 for r in results.values() if r.get('success', False))
                    cancelled = sum(1 for r in results.values() if r.get('cancelled', False))
                    total = len(results)
                    
                    if cancelled > 0:
                        st.warning(f"⚠️ Analysis stopped: {successful}/{total} completed, {cancelled} cancelled")
                    elif successful == total:
                        st.success(f"🎉 All {total} selected analyses completed successfully!")
                    else:
                        st.warning(f"⚠️ {successful}/{total} analyses completed successfully")
                    
                    # Clean up
                    st.session_state.selective_analysis_running = False
                    if 'selective_analysis_cancelled' in st.session_state:
                        del st.session_state.selective_analysis_cancelled
                    if 'selective_analysis_token' in st.session_state:
                        del st.session_state.selective_analysis_token
                    
                    time.sleep(2)  # Brief pause to show results
                    st.rerun()
                else:
                    # Analysis was cancelled
                    st.session_state.selective_analysis_running = False
                    if 'selective_analysis_cancelled' in st.session_state:
                        del st.session_state.selective_analysis_cancelled
                    if 'selective_analysis_token' in st.session_state:
                        del st.session_state.selective_analysis_token
                    st.rerun()
            
            # Separator
            st.markdown("---")
            
            # Check if full analysis is running
            full_running = st.session_state.get('full_analysis_running', False)
            
            if not full_running:
                if st.button("🚀 Run All Analyses", key="parallel_ai"):
                    # Mark as running
                    st.session_state.full_analysis_running = True
                    st.session_state.full_analysis_token = f"full_{int(time.time())}"
                    st.rerun()
            else:
                # Show progress and stop button (no columns in sidebar)
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                if st.button("🛑 Stop All Analysis", key="stop_full"):
                    st.session_state.full_analysis_cancelled = True
                    st.warning("Stopping analysis...")
                    st.rerun()
                
                # Run the analysis
                if not st.session_state.get('full_analysis_cancelled', False):
                    parallel_analyzer = ParallelAIAnalyzer(st.session_state.repo_path)
                    
                    # Create cancellation token
                    from analyzers.base_analyzer import CancellationToken
                    token = CancellationToken(st.session_state.full_analysis_token)
                    parallel_analyzer.set_cancellation_token(token)
                    
                    # Progress callback
                    def update_progress(completed, total, current_analyzer):
                        progress = (completed / total) * 100
                        progress_bar.progress(progress / 100.0)
                        status_text.text(f"Analyzing {current_analyzer.replace('_', ' ').title()}... ({completed}/{total})")
                    
                    # Run parallel analysis
                    start_time = time.time()
                    results = parallel_analyzer.run_parallel_analysis(progress_callback=update_progress)
                    end_time = time.time()
                    
                    # Update progress
                    progress_bar.progress(1.0)
                    status_text.text(f"Completed in {end_time - start_time:.1f} seconds")
                    
                    # Store results in session state
                    st.session_state.parallel_ai_results = results
                    
                    # Show summary
                    successful = sum(1 for r in results.values() if r.get('success', False))
                    cancelled = sum(1 for r in results.values() if r.get('cancelled', False))
                    total = len(results)
                    
                    if cancelled > 0:
                        st.warning(f"⚠️ Analysis stopped: {successful}/{total} completed, {cancelled} cancelled")
                    elif successful == total:
                        st.success(f"🎉 All {total} AI analyses completed successfully!")
                    else:
                        st.warning(f"⚠️ {successful}/{total} analyses completed successfully")
                    
                    # Clean up
                    st.session_state.full_analysis_running = False
                    if 'full_analysis_cancelled' in st.session_state:
                        del st.session_state.full_analysis_cancelled
                    if 'full_analysis_token' in st.session_state:
                        del st.session_state.full_analysis_token
                    
                    time.sleep(2)  # Brief pause to show results
                    st.rerun()
                else:
                    # Analysis was cancelled
                    st.session_state.full_analysis_running = False
                    if 'full_analysis_cancelled' in st.session_state:
                        del st.session_state.full_analysis_cancelled
                    if 'full_analysis_token' in st.session_state:
                        del st.session_state.full_analysis_token
                    st.rerun()
            
            # Clear AI results button
            if 'parallel_ai_results' in st.session_state:
                if st.button("🗑️ Clear AI Results", key="clear_ai"):
                    del st.session_state.parallel_ai_results
                    st.rerun()
    
    # Main content area
    if 'repo_path' in st.session_state:
        repo_path = st.session_state.repo_path
        
        # Get selected analyses from session state
        selected_analyses = st.session_state.get('selected_analyses', [])
        
        # Show Save & Export Options when analysis results are available
        if 'analysis_results' in st.session_state and selected_analyses:
            results = st.session_state.analysis_results
            successful_results = {k: v for k, v in results.items() if v.get('success', False)}
            
            if successful_results:
                st.header("💾 Save & Export Options")
                st.markdown("Download your analysis report in different formats")
                
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    # JSON Export
                    json_data = {
                        'repository': repo_path,
                        'generated': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                        'analysis_results': successful_results
                    }
                    json_str = json.dumps(json_data, indent=2, default=str)
                    st.download_button(
                        label="📄 Save as JSON",
                        data=json_str,
                        file_name=f"codebase_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                        mime="application/json",
                        help="Download analysis results as JSON file"
                    )
                
                with col2:
                    # PDF Export
                    try:
                        pdf_buffer = generate_pdf_report(successful_results, repo_path)
                        st.download_button(
                            label="📄 Download PDF Report",
                            data=pdf_buffer.getvalue(),
                            file_name=f"codebase_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                            mime="application/pdf",
                            help="Download comprehensive analysis report as PDF",
                            type="primary"
                        )
                    except ImportError as e:
                        st.error("PDF generation requires additional packages. Please install reportlab: `pip install reportlab`")
                    except Exception as e:
                        st.error(f"Error generating PDF: {str(e)}")
                
                with col3:
                    # DOCX Export  
                    try:
                        docx_buffer = generate_docx_report(successful_results, repo_path)
                        st.download_button(
                            label="📄 Download DOCX Report",
                            data=docx_buffer.getvalue(),
                            file_name=f"codebase_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            help="Download comprehensive analysis report as Word document"
                        )
                    except ImportError as e:
                        st.error("DOCX generation requires additional packages. Please install python-docx: `pip install python-docx`")
                    except Exception as e:
                        st.error(f"Error generating DOCX: {str(e)}")
                
                st.markdown("---")
        
        if selected_analyses:
            # Create tabs only for selected analyses
            analysis_options = {
                'expertise': '👥 Commits Mapping',
                'timeline': '📅 Timeline Analysis',
                'api_contracts': '🔌 API Analysis',
                'ai_context': '🤖 AI Integration Analysis',
                'risk_analysis': '⚠️ Risk Analysis',
                'development_patterns': '🏗️ Development Patterns',
                'version_governance': '📦 Version Governance',
                'tech_debt': '🔧 Tech Debt Detection',
                'design_patterns': '📐 Design Patterns'
            }
            
            # Create tabs for selected analyses
            tab_labels = [analysis_options[analysis] for analysis in selected_analyses]
            tabs = st.tabs(tab_labels)
            
            # Render each selected analyzer
            analyzer_classes = {
                'expertise': ExpertiseMapper,
                'timeline': TimelineAnalyzer,
                'api_contracts': APIContractAnalyzer,
                'ai_context': AIContextAnalyzer,
                'risk_analysis': RiskAnalysisAnalyzer
            }
            
            for i, analysis_key in enumerate(selected_analyses):
                with tabs[i]:
                    # Show individual analyzer status at the top of each tab
                    if 'analyzer_progress' in st.session_state and analysis_key in st.session_state.analyzer_progress:
                        analyzer_info = st.session_state.analyzer_progress[analysis_key]
                        status = analyzer_info.get('status', 'Not Started')
                        progress = analyzer_info.get('progress', 0)
                        
                        # Create status indicator at top of tab
                        status_container = st.container()
                        with status_container:
                            if status == 'Completed':
                                st.success(f"✅ Analysis Complete")
                            elif status == 'Running':
                                st.info(f"🔄 Analysis in Progress...")
                                if progress > 0:
                                    st.progress(progress / 100.0)
                            elif status == 'Failed':
                                st.error(f"❌ Analysis Failed")
                            elif status == 'Queued':
                                st.warning(f"⏳ Analysis Queued")
                            else:
                                st.info(f"📋 Ready to Analyze")
                        
                        st.markdown("---")
                    
                    # Show analysis results if available
                    if 'analysis_results' in st.session_state and analysis_key in st.session_state.analysis_results:
                        result = st.session_state.analysis_results[analysis_key]
                        if result.get('success', False):
                            # Render the analyzer with its data
                            analyzer_class = analyzer_classes[analysis_key]
                            analyzer = analyzer_class(repo_path)
                            analyzer.render()
                            
                            # Show AI insights
                            st.markdown("---")
                            st.subheader("🤖 AI Insights")
                            st.markdown(result['insight'])
                        else:
                            # Show error state
                            st.error(f"Analysis failed: {result.get('error', 'Unknown error')}")
                            st.info("Please try running the analysis again.")
                    else:
                        # Check if this analyzer is currently running in parallel analysis
                        if ('analyzer_progress' in st.session_state and 
                            analysis_key in st.session_state.analyzer_progress and 
                            st.session_state.analyzer_progress[analysis_key].get('status') == 'Running'):
                            
                            # Show loading message for this specific analyzer
                            analyzer_class = analyzer_classes[analysis_key]
                            analyzer = analyzer_class(repo_path)
                            
                            # Create a loading message specific to this analyzer
                            loading_messages = {
                                'expertise': "Analyzing team expertise and knowledge distribution...",
                                'timeline': "Analyzing project timeline and development patterns...",
                                'api_contracts': "Analyzing API contracts and integration points...",
                                'ai_context': "Analyzing codebase context for AI integration...",
                                'risk_analysis': "Analyzing risks and test coverage...",
                                'development_patterns': "Analyzing development patterns and code quality...",
                                'version_governance': "Analyzing version governance and dependencies...",
                                'tech_debt': "Analyzing technical debt and code complexity...",
                                'design_patterns': "Analyzing design patterns and architecture..."
                            }
                            
                            loading_message = loading_messages.get(analysis_key, f"Analyzing {analysis_key.replace('_', ' ')}...")
                            
                            # Show loading spinner and message
                            with st.spinner(loading_message):
                                st.info(f"🔄 {loading_message}")
                            
                            # Still render the analyzer interface in a disabled state
                            st.markdown("*Analysis in progress... Results will appear here when complete.*")
                            
                        else:
                            # Show placeholder when no analysis has been run yet
                            analyzer_class = analyzer_classes[analysis_key]
                            analyzer = analyzer_class(repo_path)
                            
                            # Check if this analyzer is currently selected for analysis
                            if analysis_key in st.session_state.get('selected_analyses', []):
                                st.info("👈 Click 'Run Selected Analyses' in the sidebar to analyze this repository.")
                            
                            # Still render the analyzer interface for manual use
                            analyzer.render()
        else:
            # Show guidance when no analyses are selected
            st.info("👈 Please select analyses to run from the sidebar to see the results here.")
            
            # Show available analysis types
            st.markdown("### Available Analysis Types:")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.markdown("""
                **👥 Commits Mapping**
                - Identify key contributors
                - Knowledge distribution
                - Team collaboration patterns
                
                **📅 Timeline Analysis**
                - Project activity trends
                - Development velocity
                - Commit patterns
                
                **🔌 API Analysis**
                - API endpoint analysis
                - Integration points
                - Contract validation
                """)
            
            with col2:
                st.markdown("""
                **🤖 AI Integration Analysis**
                - AI integration opportunities
                - Code structure analysis
                - Implementation recommendations
                
                **⚠️ Risk Analysis**
                - Security vulnerabilities
                - Code quality risks
                - Dependency issues
                """)
    else:
        # Show welcome screen when no repository is selected
        st.markdown("### 🚀 Welcome to AI Codebase Analyzer")
        st.markdown("""
        This tool helps you quickly understand and analyze any Git repository using AI-powered insights.
        
        **Getting Started:**
        1. 👈 Enter your repository path in the sidebar
        2. 🎯 Select the analyses you want to run
        3. 🚀 Click "Run Selected Analyses" to get AI insights
        4. 📊 View results in the tabs that appear
        
        **What you'll get:**
        - Deep insights into commit mapping with developer worked
        - Timeline analysis of project activity
        - Risk Analysis & Test Coverage
        - API Analysis & Integration Points
        - AI Integration Analysis for New Features
        """)
        
        # Show example repository paths
        st.markdown("### � Example Repository Paths:")
        st.code("""
        # Local repositories
        C:\\Projects\\my-project
        /home/user/projects/my-app
        
        # Current directory
        .
        
        # Relative paths
        ../other-project
        ./subfolder/project
        """)

if __name__ == "__main__":
    main()
